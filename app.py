import asyncio
import io
import os
import re
import socket
import time
from datetime import datetime

import cv2
from docx import Document
from dotenv import load_dotenv
import edge_tts
from flask import Flask, jsonify, render_template, request, send_file

# Google Gemini Official SDK
from google import genai
from google.genai import types
import numpy as np
from PIL import Image
import pytesseract
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

load_dotenv()

app = Flask(__name__)

# ==========================================
# GOOGLE GEMINI API CONFIGURATION & KEY ROTATION
# ==========================================
# Google Accounts ၅ ခုမှ API Key များကို Load လုပ်ခြင်း
RAW_KEYS = [
    os.environ.get("GEMINI_API_KEY_1", ""),
    os.environ.get("GEMINI_API_KEY_2", ""),
    os.environ.get("GEMINI_API_KEY_3", ""),
    os.environ.get("GEMINI_API_KEY_4", ""),
    os.environ.get("GEMINI_API_KEY_5", ""),
    os.environ.get("GEMINI_API_KEY", ""),  # Fallback Single Key
]

# လွတ်နေသော Key များနှင့် ထပ်နေသော Key များကို စစ်ထုတ်ခြင်း
API_KEYS = list(dict.fromkeys([k.strip() for k in RAW_KEYS if k and k.strip()]))

AVAILABLE_MODELS = [
    "gemini-3.7-flash",
    "gemini-3.5-flash",
    "gemini-2.5-flash",
]


# ==========================================
# INTERNET / VPN CONNECTIVITY CHECKER
# ==========================================
def is_connected(host="8.8.8.8", port=53, timeout=3):
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error:
        return False


# ==========================================
# MULTI-KEY, MULTI-MODEL FALLBACK & RETRY MECHANISM
# ==========================================
def call_gemini_with_backoff(
    contents,
    system_instruction="You are a helpful and intelligent AI assistant.",
):
    if not API_KEYS:
        raise Exception("Google Gemini API Key များကို .env ဖိုင်တွင် မတွေ့ရှိပါ၊")

    # Rate Throttling: API မခေါ်မီ Request နှုန်း ထိန်းရန် ၂ စက္ကန့် စောင့်ပေးခြင်း
    time.sleep(2)

    last_error = ""

    # 1. Key Rotation: API Key ၅ ခုကို တစ်ခုပြီးတစ်ခု လှည့်သုံးမည်
    for key_idx, key in enumerate(API_KEYS, start=1):
        try:
            client = genai.Client(api_key=key)
        except Exception as client_err:
            print(f"⚠️ Key #{key_idx} Client Initialize မရပါ: {client_err}")
            continue

        # 2. Model Fallback: Key တစ်ခုစီအတွက် Model များကို စမ်းသပ်မည်
        for model_name in AVAILABLE_MODELS:
            delay = 2
            # 3. Exponential Backoff Retry: Error တက်ပါက Delay ကို ၂ ဆ တိုး၍ Retry လုပ်မည်
            for attempt in range(3):
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=contents,
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction, temperature=0.1
                        ),
                    )
                    print(f"✅ Successfully used Key #{key_idx} with Model: {model_name}")
                    return response

                except Exception as e:
                    err_str = str(e)
                    last_error = err_str

                    if (
                        "RESOURCE_EXHAUSTED" in err_str
                        or "429" in err_str
                        or "Quota" in err_str
                    ):
                        print(
                            f"⚠️ Key #{key_idx} ({model_name}) Quota/Limit ပြည့်သွားပါပြီ။ "
                            "နောက် Key သို့ ပြောင်းနေပါသည်..."
                        )
                        # Quota ပြည့်ပါက ယခု Key ကို ကျော်ပြီး နောက် Key တစ်ခုသို့ ချက်ချင်း သွားမည်
                        break

                    elif "404" in err_str:
                        print(f"⚠️ Model '{model_name}' မတွေ့ပါ။ နောက် Model သို့ ပြောင်းပါမည်...")
                        break

                    elif any(
                        x in err_str
                        for x in [
                            "503",
                            "UNAVAILABLE",
                            "disconnected",
                            "high demand",
                            "timed out",
                        ]
                    ):
                        print(
                            f"⚠️ Key #{key_idx} [{model_name}] Server Busy ဖြစ်နေ၍ {delay} စက္ကန့်"
                            f" စောင့်ပါသည်... (ကြိုးစားမှု - {attempt+1})"
                        )
                        time.sleep(delay)
                        delay *= 2  # Exponential Backoff (2s -> 4s -> 8s)
                    else:
                        print(f"⚠️ Key #{key_idx} [{model_name}] Error: {err_str}")
                        break

    raise Exception(
        f"API Keys အားလုံး သို့မဟုတ် Quota အားလုံး ကုန်ဆုံးသွားပါပြီ။ (နောက်ဆုံး Error: {last_error})"
    )


def preprocess_image_for_gemini(image_bytes):
    file_bytes = np.frombuffer(image_bytes, np.uint8)
    image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError("Invalid image format or corrupted file.")
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    return Image.fromarray(image_rgb)


# ==========================================
# HIGH ACCURACY LOCAL TESSERACT PREPROCESSING
# ==========================================
def preprocess_image_for_tesseract(image_bytes):
    file_bytes = np.frombuffer(image_bytes, np.uint8)
    image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

    if image is None:
        raise ValueError("Invalid image format or corrupted file.")

    height, width = image.shape[:2]
    if width < 2200:
        scale = 2200 / width
        image = cv2.resize(
            image,
            (int(width * scale), int(height * scale)),
            interpolation=cv2.INTER_CUBIC,
        )

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)
    denoised = cv2.GaussianBlur(enhanced, (3, 3), 0)

    return Image.fromarray(denoised)


# ==========================================
# FIXED OCR CLEANER
# ==========================================
def clean_tesseract_myanmar_text(text, is_myanmar_selected=True):
    if not text:
        return ""

    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        line_clean = re.sub(
            r"\[Geden\]?|[°¢\[\]\{\}\#\/\=\_\<\>]", "", line_str
        )
        line_clean = re.sub(r"\s+", " ", line_clean).strip()

        if is_myanmar_selected:
            line_clean = line_clean.replace("!", "။").replace("|", "၊")

        cleaned_no_punct = re.sub(
            r"[\(\)\,\.\;\:\-\"\'\?\!\၊\။]", "", line_clean
        ).strip()
        myanmar_chars = re.findall(r"[\u1000-\u109F]", cleaned_no_punct)
        latin_chars = re.findall(r"[a-zA-Z0-9]", cleaned_no_punct)
        tokens = cleaned_no_punct.split()

        if (
            len(myanmar_chars) <= 6
            and tokens
            and (len(myanmar_chars) / len(tokens)) <= 2.2
        ):
            if len(latin_chars) < 3:
                continue

        if (
            re.match(r"^(?:[\u1000\u101d\u1040-\u1049\s\(\)]){3,}$", line_clean)
            and len(latin_chars) == 0
        ):
            continue

        if line_clean:
            cleaned_lines.append(line_clean)

    return "\n".join(cleaned_lines)


@app.route("/")
def index():
    return render_template("index.html")


# ==========================================
# OCR SCAN ENDPOINT
# ==========================================
@app.route("/ocr", methods=["POST"])
@app.route("/scan", methods=["POST"])
def scan_ocr():
    file = request.files.get("file") or request.files.get("image")
    lang_choice = request.form.get("language", "mya+eng")

    if not file or file.filename == "":
        return jsonify({"success": False, "error": "No image file uploaded"}), 400

    image_bytes = file.read()
    extracted_text = ""

    if not is_connected():
        try:
            processed_img = preprocess_image_for_tesseract(image_bytes)
            custom_config = r"--oem 3 --psm 4"
            raw_text = pytesseract.image_to_string(
                processed_img, lang=lang_choice, config=custom_config
            ).strip()

            extracted_text = clean_tesseract_myanmar_text(
                raw_text, is_myanmar_selected=("mya" in lang_choice)
            )
        except Exception as t_err:
            return (
                jsonify({
                    "success": False,
                    "error": f"Local OCR Extraction Failed: {str(t_err)}",
                }),
                500,
            )

    else:
        try:
            pil_img = preprocess_image_for_gemini(image_bytes)
            ocr_prompt = (
                "Extract all text from this image with absolute accuracy. "
                "Pay close attention to Myanmar script, English text, tables, and"
                " numbers. "
                "Preserve the original layout and line breaks. "
                "Provide ONLY the extracted text without any introductory remarks."
            )
            sys_inst = (
                "You are an expert OCR engine specialized in extracting text with"
                " high accuracy."
            )

            response = call_gemini_with_backoff(
                [pil_img, ocr_prompt], system_instruction=sys_inst
            )
            extracted_text = response.text.strip()

        except Exception as e:
            try:
                processed_img = preprocess_image_for_tesseract(image_bytes)
                custom_config = r"--oem 3 --psm 4"
                raw_text = pytesseract.image_to_string(
                    processed_img, lang=lang_choice, config=custom_config
                ).strip()

                extracted_text = clean_tesseract_myanmar_text(
                    raw_text, is_myanmar_selected=("mya" in lang_choice)
                )
            except Exception as t_err:
                return (
                    jsonify(
                        {"success": False, "error": f"OCR Extraction Failed: {str(t_err)}"}
                    ),
                    500,
                )

    return jsonify({
        "success": True,
        "text": extracted_text,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    })


# ==========================================
# GOOGLE GEMINI AI ASSISTANT ENDPOINT
# ==========================================
@app.route("/ai/process", methods=["POST"])
def ai_process():
    data = request.get_json() or {}
    text = data.get("text", "")
    action = data.get("action", "")
    target_lang = data.get("target_lang", "")
    user_prompt = data.get("prompt") or data.get("query", "")

    if not is_connected():
        return (
            jsonify({
                "success": False,
                "error": "AI features require an active internet connection or VPN.",
            }),
            400,
        )

    if not text and action != "ask":
        return (
            jsonify(
                {"success": False, "error": "No text provided for AI processing"}
            ),
            400,
        )

    try:
        if action == "translate":
            if target_lang == "en":
                lang_instruction = (
                    "Translate the input text into clear and standard English."
                )
            elif target_lang == "my":
                lang_instruction = (
                    "Translate the input text into natural Myanmar (Burmese) language."
                )
            else:
                lang_instruction = (
                    "If the input text is primarily in Myanmar, translate it to English."
                    " If English, translate it to Myanmar."
                )

            prompt = (
                f"You are a professional translator. {lang_instruction}\nPreserve"
                " document structure and formatting. Provide ONLY the final"
                f" translation.\n\nInput Text:\n{text}"
            )

        elif action == "summarize":
            prompt = (
                "Please provide a concise, well-structured bullet-point summary of"
                f" the following text:\n\n{text}"
            )

        elif action in ["fix", "spellcheck"]:
            prompt = (
                "Please review and fix all spelling errors, grammar mistakes, and"
                " typos in the following text. Provide ONLY the corrected"
                f" text:\n\n{text}"
            )

        elif action == "explain":
            prompt = user_prompt or f"Explain the main points of:\n\n{text}"

        elif action == "ask":
            prompt = f"Context:\n{text}\n\nQuestion: {user_prompt}"

        else:
            return jsonify({"success": False, "error": "Invalid AI action"}), 400

        sys_inst = "You are a helpful AI assistant specialized in text processing."
        response = call_gemini_with_backoff(prompt, system_instruction=sys_inst)

        return jsonify({"success": True, "result": response.text.strip()})

    except Exception as e:
        return (
            jsonify({"success": False, "error": f"Gemini AI Error: {str(e)}"}),
            500,
        )


# ==========================================
# EXPORT ROUTES (TXT, DOCX & PDF)
# ==========================================
@app.route("/export/txt", methods=["POST"])
def export_txt():
    data = request.get_json() or {}
    text = data.get("text", "")
    byte_io = io.BytesIO(text.encode("utf-8"))
    return send_file(
        byte_io,
        mimetype="text/plain;charset=utf-8",
        as_attachment=True,
        download_name="Extracted_Text.txt",
    )


@app.route("/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json() or {}
    text = data.get("text", "")

    doc = Document()
    doc.add_heading("OCR Extracted Document", level=1)
    for line in text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(
        doc_io,
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        as_attachment=True,
        download_name="ocr_result.docx",
    )


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json() or {}
    text = data.get("text", "")

    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=letter)
    styles = getSampleStyleSheet()

    story = [Paragraph("OCR Extracted Document", styles["Heading1"]), Spacer(1, 18)]

    for line in text.split("\n"):
        if line.strip():
            clean_line = (
                line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            story.append(Paragraph(clean_line, styles["Normal"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    pdf_io.seek(0)

    return send_file(
        pdf_io,
        mimetype="application/pdf",
        as_attachment=True,
        download_name="result.pdf",
    )


# ==========================================
# FAST EDGE-TTS ENDPOINT FOR FLASK
# ==========================================
@app.route("/api/tts", methods=["POST"])
def tts_endpoint():
    data = request.get_json() or {}
    text = data.get("text", "").strip()
    lang = data.get("lang", "en")

    if not text:
        return jsonify({"error": "No text provided"}), 400

    voice = "my-MM-NilarNeural" if lang == "my" else "en-US-AvaNeural"

    async def generate_audio():
        communicate = edge_tts.Communicate(text, voice)
        audio_bytes = b""
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_bytes += chunk["data"]
        return audio_bytes

    try:
        audio_data = asyncio.run(generate_audio())
        return send_file(
            io.BytesIO(audio_data),
            mimetype="audio/mpeg",
            as_attachment=False,
            download_name="speech.mp3",
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)